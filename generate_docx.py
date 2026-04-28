# -*- coding: utf-8 -*-
"""
457DEEP 공기업 면접 강의 시리즈 - 8편: 관심 사업
SKILL.md 규칙에 따라 docx 생성
"""

import subprocess
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "8편_우리_공기업_사업_중_관심있는_것과_이유.docx"


def git(*args, check=True):
    return subprocess.run(
        ["git", "-C", str(ROOT), *args],
        check=check,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def push_to_github(file_path: Path, topic: str):
    rel = file_path.relative_to(ROOT).as_posix()
    git("add", rel)
    status = git("status", "--porcelain")
    if not status.stdout.strip():
        print("no changes to commit")
        return
    msg = f"add output: {topic}"
    git("commit", "-m", msg)
    result = git("push", "origin", "main", check=False)
    if result.returncode != 0:
        print("push failed:", result.stderr, file=sys.stderr)
        sys.exit(1)
    print("pushed to origin/main")

FONT_KR = "맑은 고딕"
FONT_EN = "Arial"


def set_run_font(run, size_pt, bold=False, color=None):
    run.font.name = FONT_EN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)
    rfonts.set(qn("w:cs"), FONT_EN)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(p, line=1.5, before=0, after=6):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_paragraph(doc, text, size=11, bold=False, align=None, indent_left=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    set_paragraph_spacing(p, line=1.5, after=6)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=1.5, before=12, after=12)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line=1.5, before=14, after=8)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    return p


def add_blockquote(doc, lines):
    """인용블록 처리: 좌측 들여쓰기 + 회색 배경"""
    add_blank(doc)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        set_paragraph_spacing(p, line=1.5, before=0, after=4)
        # 좌측 보더 추가 (인용 표시)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "10")
        left.set(qn("w:color"), "888888")
        pBdr.append(left)
        pPr.append(pBdr)
        # 셀 배경 (회색)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        run = p.add_run(line)
        set_run_font(run, 11)
    add_blank(doc)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    set_paragraph_spacing(p, line=1.5, after=4)
    run = p.add_run(f"• {text}")
    set_run_font(run, 11)
    return p


def add_number(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    set_paragraph_spacing(p, line=1.5, after=4)
    run = p.add_run(f"{num}. {text}")
    set_run_font(run, 11)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line=1.0, after=0)
    return p


def add_divider(doc):
    """섹션 구분: 빈줄 + --- + 빈줄"""
    add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=1.0, after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)
    add_blank(doc)


def build_document():
    doc = Document()

    # 페이지 설정 (A4, 여백 2.54cm)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 기본 스타일
    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)

    # ===== 표지 =====
    add_h1(doc, "457DEEP 공기업 면접 강의 시리즈")
    add_h1(doc, "8편 - 우리 공기업 사업 중에 관심있는 것과 이유는?")
    add_blank(doc)

    # ===== 시작 고정 문구 =====
    add_paragraph(
        doc,
        "▶️알려드리는 말씀 : 현재 공기업 면접 실전 강의안을 전면 개편하여 업로드 중에 있습니다. "
        "최대한 빠른 시일 내, 계속 업로드 될 예정입니다. 이 강의를 보기 전에, 면접 기본 강의도 꼭 참고해주세요 :) "
        "특히 10D 공식 1-5편은 해당 공기업 면접 준비 방법의 기본이 되는 공식이니 꼭 숙지하시기 바랍니다.",
        size=11,
    )
    add_paragraph(doc, "https://457deep.com/beginner/interview", size=11, color=(0, 102, 204))

    add_divider(doc)

    # ===== 도입부 =====
    add_h3(doc, "🎯 이 편에서 다룰 내용")
    add_paragraph(
        doc,
        "이번 편에서는 \"우리 공기업 사업 중에 관심있는 것과 이유는?\" 질문을 다룹니다. "
        "면접관이 이 질문에서 무엇을 보는지, 그리고 어떤 흐름으로 답을 설계해야 합격 답변이 되는지 함께 정리해 보시기 바랍니다.",
    )

    add_divider(doc)

    # ===== 1. 면접관 의도 =====
    add_h3(doc, "🔍 면접관이 이 질문에서 보는 것")
    add_paragraph(
        doc,
        "이 질문은 단순한 \"좋아하는 사업 고르기\"가 아닙니다. 면접관은 다음 세 가지를 동시에 확인합니다.",
    )
    add_blank(doc)
    add_bullet(doc, "지원자가 우리 기관 사업을 실제로 공부했는가 (경영 공시, 중장기 계획, 기관장 취임사 수준까지)")
    add_bullet(doc, "관심 사업과 본인 직무, 본인 경험 사이의 연결고리가 명확한가")
    add_bullet(doc, "공익, 공정성, 사회적 가치 관점에서 사업의 의미를 해석할 수 있는가")
    add_blank(doc)
    add_paragraph(
        doc,
        "즉, \"내가 좋아하는 사업\"이 아니라 \"내가 기여할 수 있는 사업\"으로 답변의 무게 중심을 옮기셔야 합니다.",
    )

    add_divider(doc)

    # ===== 2. 답변 설계 프레임 =====
    add_h3(doc, "🧭 답변 설계 5단계 프레임")
    add_paragraph(doc, "관심 사업 답변은 다음 5단계로 설계하시기 바랍니다.")
    add_blockquote(
        doc,
        [
            "1. [두괄식] 저는 ㅇㅇ 사업에 가장 관심이 있습니다.",
            "2. [근거] 그 이유는 해당 사업이 ㅇㅇ 측면에서 공익에 직결되기 때문입니다.",
            "3. [기관 연결] 특히 우리 기관의 중장기 계획 ㅇㅇ 과제와 직접 맞닿아 있습니다.",
            "4. [본인 연결] 저는 ㅇㅇ 경험을 통해 해당 사업 수행에 보탤 수 있는 역량을 쌓아 왔습니다.",
            "5. [기여] 입사 후 ㅇㅇ 업무를 통해 해당 사업의 수혜 범위 확대에 기여하겠습니다.",
        ],
    )
    add_paragraph(
        doc,
        "이 5단계는 두괄식 → 공익 근거 → 기관 사업 연결 → 본인 경험 연결 → 기여 마무리의 흐름입니다. "
        "단계가 하나라도 빠지면 답변이 가벼워지니, 빈 칸을 본인 정보로 채우는 연습을 반드시 해두시기 바랍니다.",
    )

    add_divider(doc)

    # ===== 3. 사업 선정 기준 =====
    add_h3(doc, "📋 어떤 사업을 선택해야 하는가")
    add_paragraph(doc, "관심 사업을 고를 때는 다음 기준으로 좁혀 가시기 바랍니다.")
    add_blank(doc)
    add_number(doc, "기관의 중장기 계획에 명시된 핵심 사업 (경영 공시 기준)", 1)
    add_number(doc, "기관장 취임사 또는 신년사에서 강조된 방향성과 일치하는 사업", 2)
    add_number(doc, "본인이 지원한 직무가 실제로 수행하는 사업", 3)
    add_number(doc, "본인 경험과 연결고리가 자연스럽게 만들어지는 사업", 4)
    add_blank(doc)
    add_paragraph(
        doc,
        "여러분이 흔히 하시는 실수는 뉴스 한두 줄만 보고 사업명을 던지는 것입니다. "
        "면접관은 후속 질문으로 \"그 사업의 추진 배경은 무엇인가요?\", \"그 사업의 1차 수혜자는 누구인가요?\"를 바로 던집니다. "
        "사업 한 개를 고르되, 추진 배경과 수혜자, 사회적 의미까지 같이 준비하시기 바랍니다.",
    )

    add_divider(doc)

    # ===== 4. 답변 예시 =====
    add_h3(doc, "📚 답변 예시 (직군별)")
    add_paragraph(doc, "예시 직군별로 첫 문장 출발점과 강조점을 다르게 설계하였으니, 본인 직군에 맞춰 참고하시기 바랍니다.")

    add_blockquote(
        doc,
        [
            "예시 1) [행정직] - 가입자 보호 사업과 본인의 데이터 분석 경험 연결",
            "",
            "저는 우리 기관의 ㅇㅇ 보호 사업에 가장 관심이 있습니다. 해당 사업은 가입자 ㅇㅇ만 명의 권익에 직결되는 핵심 사업이라고 판단했습니다. "
            "특히 중장기 계획에서 가입자 정보 보호 체계 고도화가 ㅇㅇ 과제로 명시되어 있는 점이 눈에 띄었습니다. "
            "저는 학부 시절 ㅇㅇ 동아리에서 ㅇㅇ개월 간 데이터 정합성 점검 프로젝트를 수행하며, 절차와 원칙을 지키는 업무 태도를 익혀 왔습니다. "
            "입사 후에는 가입자 보호 절차 운영 업무를 정확하게 수행하여 사업의 수혜 범위 확대에 보태겠습니다.",
        ],
    )

    add_blockquote(
        doc,
        [
            "예시 2) [기술직] - 안전 관리 사업과 현장 실습 경험 연결",
            "",
            "관심 사업은 우리 기관의 ㅇㅇ 안전 관리 사업입니다. 해당 사업은 국민 ㅇㅇ만 명의 일상 안전과 직결되는 공공성이 가장 높은 영역이라고 보았습니다. "
            "기관장 취임사에서도 무사고 운영 체계 강화가 ㅇㅇ 핵심 방향으로 제시되어 있었습니다. "
            "저는 ㅇㅇ개월의 현장 실습 기간 동안 점검 매뉴얼 기반 설비 진단을 ㅇㅇ회 수행한 경험이 있습니다. "
            "입사 후에는 점검 절차 준수와 이상 징후 조기 보고를 통해 안전 관리 사업 수행에 기여하겠습니다.",
        ],
    )

    add_blockquote(
        doc,
        [
            "예시 3) [행정직] - 사회적 가치 사업과 봉사 경험 연결",
            "",
            "저의 관심 사업은 우리 기관의 사회적 가치 ㅇㅇ 사업입니다. 해당 사업은 취약 계층 수혜자 ㅇㅇ만 명에게 실질적 혜택이 닿는 사업이기 때문입니다. "
            "특히 경영 공시상 사회적 가치 지표가 ㅇㅇ% 비중으로 관리되고 있는 점에서 기관의 의지를 확인했습니다. "
            "저는 ㅇㅇ개월 간 지역 ㅇㅇ 봉사 활동에 참여하며 수혜자 관점에서 정책의 실효성을 살피는 시각을 길렀습니다. "
            "입사 후에는 사업 절차의 투명성을 지키며 수혜자 확대 업무를 수행하겠습니다.",
        ],
    )

    add_divider(doc)

    # ===== 5. 좋은 예 / 나쁜 예 =====
    add_h3(doc, "⚠️ 자주 나오는 실수 패턴")

    add_blockquote(
        doc,
        [
            "❌ 나쁜 예",
            "",
            "저는 우리 기관의 디지털 전환 사업에 관심이 많습니다. 요즘 디지털이 중요하기 때문입니다. "
            "입사하면 빠르게 적응하고 많이 배워서 성장하는 인재가 되겠습니다.",
            "",
            "→ 사업 선택 근거가 두루뭉술합니다. 기관 사업과 본인 경험 연결고리가 없습니다. "
            "마무리도 적응/배움/성장 톤이라 공기업 답변으로는 가벼운 인상을 줍니다.",
        ],
    )

    add_blockquote(
        doc,
        [
            "⭕️ 좋은 예",
            "",
            "저는 우리 기관의 ㅇㅇ 디지털 전환 사업에 가장 관심이 있습니다. 이 사업은 가입자 ㅇㅇ만 명의 행정 편의와 직결되며, "
            "중장기 계획상 ㅇㅇ 과제로 명시되어 있습니다. 저는 ㅇㅇ개월 간 ㅇㅇ 데이터 표준화 프로젝트를 수행한 경험이 있습니다. "
            "입사 후에는 절차와 원칙을 지키며 디지털 전환 업무를 정확하게 수행하여, 가입자 편의 확대에 기여하겠습니다.",
        ],
    )

    add_divider(doc)

    # ===== 6. 체크리스트 =====
    add_h3(doc, "✅ 답변 점검 체크리스트")
    add_paragraph(doc, "본인 답변을 점검할 때 다음 항목을 반드시 확인하시기 바랍니다.")
    add_blank(doc)
    add_bullet(doc, "두괄식으로 사업명을 먼저 던졌는가")
    add_bullet(doc, "사업 선택 근거가 공익, 공정성, 사회적 가치 관점에서 설명되는가")
    add_bullet(doc, "기관 중장기 계획 또는 경영 공시상 근거가 포함되어 있는가")
    add_bullet(doc, "본인 경험과의 연결고리가 자연스럽게 이어지는가")
    add_bullet(doc, "마무리가 기여, 보태기, 수행 톤인가 (적응, 배움, 성장 마무리 아님)")
    add_bullet(doc, "사기업 용어 (매출, 경쟁력, 시장 점유율 등) 가 들어가지 않았는가")
    add_blank(doc)

    add_divider(doc)

    # ===== 7. 마무리 =====
    add_h3(doc, "📝 이 편 핵심 요약")
    add_paragraph(
        doc,
        "관심 사업 질문은 \"내가 좋아하는 사업 말하기\"가 아니라 \"내가 기여할 사업의 근거 대기\"입니다. "
        "5단계 프레임 - 두괄식, 공익 근거, 기관 사업 연결, 본인 경험 연결, 기여 마무리 - 를 본인 정보로 채워 보시기 바랍니다.",
    )
    add_blank(doc)
    add_paragraph(
        doc,
        "다음 편에서는 9편 - 위기와 기회 답변 설계법을 다룰 예정입니다. "
        "기관의 위기 요인을 어떻게 진단하고, 그 안에서 기회를 어떻게 제시할지 함께 정리해 보시기 바랍니다.",
    )

    add_divider(doc)

    # ===== 끝 고정 문구 =====
    add_paragraph(doc, "면접 답변 첨삭은 아래 1대 1 피드백/ 첨삭으로 신청해주세요")
    add_paragraph(doc, "https://457deep.com/feedbacks", color=(0, 102, 204))

    return doc


if __name__ == "__main__":
    doc = build_document()
    doc.save(str(OUTPUT_PATH))
    print(f"saved: {OUTPUT_PATH}")

    no_push = "--no-push" in sys.argv
    if not no_push:
        push_to_github(OUTPUT_PATH, "8편 - 관심 사업")
