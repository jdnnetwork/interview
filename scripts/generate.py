"""457DEEP 공기업 면접 강의 시리즈 자동 생성 스크립트.

흐름:
- inputs/input.md (주제 + 주문 사항) 와 SKILL.md (시스템 규칙) 읽음
- Anthropic Claude API (claude-opus-4-7, adaptive thinking) 2회 호출
  - 1차: 1편 본편 마크다운 + SHORT_TITLE
  - 2차: 1편 본문을 컨텍스트로 넣어 2편 본편 작성 (이어지는 심화 내용)
- python-docx 로 각각 docx 렌더링
- outputs/YYYY-MM-DD_<제목>_1편.docx, ..._2편.docx 저장 (KST 기준 날짜)
"""

from __future__ import annotations

import os
import re
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
INPUT_PATH = ROOT / "inputs" / "input.md"
SKILL_PATH = ROOT / "SKILL.md"
OUTPUT_DIR = ROOT / "outputs"

KST = timezone(timedelta(hours=9))
FONT_KR = "맑은 고딕"
FONT_EN = "Arial"

MODEL = "claude-opus-4-7"
MAX_TOKENS = 32000

USER_FORMAT_SUFFIX_PART1 = """

---

[출력 형식 - 반드시 지킬 것]

이번 강의는 "1편 + 2편" 시리즈로 작성합니다. 지금은 **1편**을 작성합니다.

첫 번째 줄: SHORT_TITLE: <4~10자 짧은 한국어 제목 (파일명용, 예: 관심사업, 자기소개, 직무역량)>
두 번째 줄: ---
이후: SKILL.md 의 모든 규칙을 따른 강의 1편 본편 마크다운만 작성. 다른 설명/메타텍스트/code fence 금지.
시작 고정 문구와 끝 고정 문구를 반드시 포함 (1편은 독립된 docx 파일).

[분량 및 구성 가이드]
- 1편 분량은 충분히 길게 (한국어 약 8,000자 이상, 가능하면 10,000자 내외)
- 도입 ~ 핵심 개념 ~ 기본 예시까지 다루고, 더 심화된 사례/실전 응용은 2편으로 자연스럽게 넘기는 흐름으로 마무리
"""

USER_FORMAT_SUFFIX_PART2 = """

---

[지시사항]

위는 방금 작성한 같은 주제의 "1편" 본편입니다. 지금은 **2편**을 작성합니다.

요건:
- 1편의 흐름과 자연스럽게 연결되도록 후속 내용으로 작성
- 1편에서 이미 다룬 내용은 반복 금지. 더 심화 / 구체적 / 실전 예시 / 사례 분석 / 응용 위주
- SKILL.md 모든 규칙 그대로 적용
- 2편도 독립된 docx 파일이므로 시작 고정 문구와 끝 고정 문구 모두 포함
- 2편 분량도 충분히 길게 (한국어 약 8,000자 이상, 가능하면 10,000자 내외)

[출력 형식 - 반드시 지킬 것]

SHORT_TITLE 라인 없이 바로 본편 시작.
SKILL.md 규칙대로 강의 2편 본편 마크다운만. 다른 설명/메타텍스트/code fence 금지.
"""


# ===== Claude API =====

def call_claude(skill_md: str, input_md: str, prior_part: str | None = None) -> str:
    """prior_part=None 이면 1편 호출, 주어지면 그 본편을 컨텍스트로 2편 호출."""
    client = anthropic.Anthropic()

    if prior_part is None:
        user_content = input_md.strip() + USER_FORMAT_SUFFIX_PART1
        label = "part1"
    else:
        user_content = (
            "[원본 주제 및 주문 사항]\n\n"
            + input_md.strip()
            + "\n\n---\n\n[방금 작성한 1편 본편]\n\n"
            + prior_part
            + USER_FORMAT_SUFFIX_PART2
        )
        label = "part2"

    chunks: list[str] = []
    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        system=[
            {
                "type": "text",
                "text": skill_md,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[
            {
                "role": "user",
                "content": user_content,
            }
        ],
    ) as stream:
        for text in stream.text_stream:
            chunks.append(text)
        final = stream.get_final_message()
    print(
        f"usage[{label}]: input={final.usage.input_tokens}, "
        f"output={final.usage.output_tokens}, "
        f"cache_read={final.usage.cache_read_input_tokens}, "
        f"cache_create={final.usage.cache_creation_input_tokens}",
        file=sys.stderr,
    )
    return "".join(chunks)


def parse_response(text: str) -> tuple[str, str]:
    m = re.search(
        r"SHORT_TITLE\s*:\s*(.+?)\s*\n\s*-{3,}\s*\n(.+)",
        text,
        re.DOTALL,
    )
    if not m:
        return "강의", text.strip()
    return m.group(1).strip(), m.group(2).strip()


def slugify_kr(title: str) -> str:
    cleaned = re.sub(r"\s+", "", title)
    cleaned = re.sub(r"[^\w가-힣]", "", cleaned)
    return cleaned or "강의"


# ===== docx rendering =====

def _set_run_font(run, size_pt, bold=False, color=None):
    run.font.name = FONT_EN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)
    rfonts.set(qn("w:cs"), FONT_EN)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_paragraph_spacing(p, line=1.5, before=0, after=6):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_paragraph(doc, text, size=11, bold=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _set_paragraph_spacing(p, line=1.5, after=6)
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold, color=color)


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line=1.5, before=12, after=12)
    run = p.add_run(text)
    _set_run_font(run, 16, bold=True)


def _add_h3(doc, text):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=1.5, before=14, after=8)
    run = p.add_run(text)
    _set_run_font(run, 12, bold=True)


def _add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    _set_paragraph_spacing(p, line=1.5, after=4)
    run = p.add_run(f"• {text}")
    _set_run_font(run, 11)


def _add_number(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    _set_paragraph_spacing(p, line=1.5, after=4)
    run = p.add_run(f"{num}. {text}")
    _set_run_font(run, 11)


def _starts_with_emoji(text: str) -> bool:
    if not text:
        return False
    code = ord(text[0])
    return (
        0x1F000 <= code <= 0x1FFFF
        or 0x2600 <= code <= 0x27BF
        or 0x2300 <= code <= 0x23FF
        or 0x25A0 <= code <= 0x25FF
        or 0x2B00 <= code <= 0x2BFF
    )


def _is_emoji_heading(text: str) -> bool:
    if not _starts_with_emoji(text) or len(text) > 60:
        return False
    return not text.endswith((".", "!", "?", "요", "다", ":"))


def _add_blank(doc):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=1.0, after=0)


def _add_divider(doc):
    _add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line=1.0, after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _add_blank(doc)


def _add_blockquote(doc, lines):
    _add_blank(doc)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        _set_paragraph_spacing(p, line=1.5, before=0, after=4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "10")
        left.set(qn("w:color"), "888888")
        pBdr.append(left)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        run = p.add_run(line)
        _set_run_font(run, 11)
    _add_blank(doc)


def _setup_document(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)


def render_to_docx(title: str, body_md: str, output_path: Path):
    doc = Document()
    _setup_document(doc)

    _add_h1(doc, "457DEEP 공기업 면접 강의 시리즈")
    _add_h1(doc, title)
    _add_blank(doc)

    blockquote_lines: list[str] = []

    def flush_bq():
        if blockquote_lines:
            _add_blockquote(doc, list(blockquote_lines))
            blockquote_lines.clear()

    for raw_line in body_md.split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("### "):
            flush_bq()
            _add_h3(doc, stripped[4:].strip())
        elif stripped == "---":
            flush_bq()
            _add_divider(doc)
        elif stripped in ("&nbsp;", "\\&nbsp;", ""):
            flush_bq()
        elif stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            blockquote_lines.append(content)
        elif re.match(r"^\d+\.\s", stripped):
            flush_bq()
            num, text = stripped.split(".", 1)
            _add_number(doc, text.strip(), num)
        elif stripped.startswith("* ") or stripped.startswith("- "):
            flush_bq()
            _add_bullet(doc, stripped[2:].strip())
        elif _is_emoji_heading(stripped):
            flush_bq()
            _add_h3(doc, stripped)
        else:
            flush_bq()
            _add_paragraph(doc, stripped)

    flush_bq()
    doc.save(str(output_path))


# ===== entry =====

def main():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY env var not set", file=sys.stderr)
        sys.exit(1)

    if not INPUT_PATH.exists():
        print(f"ERROR: {INPUT_PATH} not found", file=sys.stderr)
        sys.exit(1)
    if not SKILL_PATH.exists():
        print(f"ERROR: {SKILL_PATH} not found", file=sys.stderr)
        sys.exit(1)

    skill_md = SKILL_PATH.read_text(encoding="utf-8")
    input_md = INPUT_PATH.read_text(encoding="utf-8")
    if not input_md.strip():
        print("ERROR: input.md is empty", file=sys.stderr)
        sys.exit(1)

    OUTPUT_DIR.mkdir(exist_ok=True)
    today = datetime.now(KST).strftime("%Y-%m-%d")

    print("calling Claude API (1편)...", file=sys.stderr)
    part1_response = call_claude(skill_md, input_md)
    title, part1_body = parse_response(part1_response)
    print(f"title: {title}", file=sys.stderr)
    print(f"part1 body length: {len(part1_body)} chars", file=sys.stderr)

    slug = slugify_kr(title)
    part1_path = OUTPUT_DIR / f"{today}_{slug}_1편.docx"
    render_to_docx(f"{title} (1편)", part1_body, part1_path)
    print(f"saved: {part1_path}", file=sys.stderr)

    print("calling Claude API (2편)...", file=sys.stderr)
    part2_response = call_claude(skill_md, input_md, prior_part=part1_body)
    # 2편은 SHORT_TITLE 없이 바로 본편이지만, 모델이 혹시 넣었으면 떼어냄
    if part2_response.strip().startswith("SHORT_TITLE"):
        _, part2_body = parse_response(part2_response)
    else:
        part2_body = part2_response.strip()
    print(f"part2 body length: {len(part2_body)} chars", file=sys.stderr)

    part2_path = OUTPUT_DIR / f"{today}_{slug}_2편.docx"
    render_to_docx(f"{title} (2편)", part2_body, part2_path)
    print(f"saved: {part2_path}", file=sys.stderr)

    print(part1_path.relative_to(ROOT).as_posix())
    print(part2_path.relative_to(ROOT).as_posix())


if __name__ == "__main__":
    main()
